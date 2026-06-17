"""
Script to reformat SECTION 11.docx into a well-structured document
with proper headings for Table of Contents generation.
Version 2 - Improved structure detection
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_toc(doc):
    """Add a Table of Contents field to the document"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "[Right-click and select 'Update Field' to generate TOC]"
    
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    r_element.append(fldChar4)
    
    return paragraph

def is_section_header(text):
    """Check if text is a main section header like 'SECTION 3' or 'SECTION 4'"""
    pattern = r'^SECTION\s*\d+'
    return bool(re.match(pattern, text.upper().strip()))

def is_subsection_header(text):
    """Check if text is a subsection like '3.1', '3.2', '4.1'"""
    pattern = r'^\d+\.\d+\s'
    return bool(re.match(pattern, text.strip()))

def is_question_header(text):
    """Check if text is a question like 'Q3.1.1:', 'Q4.2.3:'"""
    pattern = r'^Q\d+\.\d+\.?\d*:'
    return bool(re.match(pattern, text.strip()))

def is_month_header(text):
    """Check if text is a month header"""
    return text.lower().strip().startswith('month') or 'شهر' in text

def is_week_header(text):
    """Check if text is a week header"""
    text_lower = text.lower().strip()
    return text_lower.startswith('week') or 'أسبوع' in text or 'الأسبوع' in text

def is_module_header(text):
    """Check if text is a module header like '3.3.1 Employee Request Forms'"""
    pattern = r'^\d+\.\d+\.\d+\s'
    return bool(re.match(pattern, text.strip()))

def is_category_header(text):
    """Check if text is a task category"""
    categories = ['Backend', 'Frontend', 'Database', 'System Analysis', 
                  'Maintenance', 'Quality Assurance', 'Integration', 
                  'Task Categories']
    for cat in categories:
        if text.strip().startswith(cat):
            return True
    return False

def main():
    # Paths
    input_path = r"C:\Users\sqlcc\OneDrive\Desktop\THE COPY\project-root_server_v\SECTION 11.docx"
    output_path = r"C:\Users\sqlcc\OneDrive\Desktop\THE COPY\project-root_server_v\SECTION 11 - Formatted.docx"
    
    print(f"Reading document: {input_path}")
    original_doc = Document(input_path)
    
    # Create new document
    new_doc = Document()
    
    # Configure styles
    styles = new_doc.styles
    
    # Title style
    title_style = styles['Title']
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(26)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 51, 102)
    
    # Heading 1 - Main Sections
    h1_style = styles['Heading 1']
    h1_style.font.name = 'Arial'
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 51, 102)
    
    # Heading 2 - Subsections / Months / Weeks
    h2_style = styles['Heading 2']
    h2_style.font.name = 'Arial'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 76, 153)
    
    # Heading 3 - Questions / Modules
    h3_style = styles['Heading 3']
    h3_style.font.name = 'Arial'
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(51, 51, 51)
    
    # Normal style
    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    
    # Extract all paragraphs
    all_paragraphs = []
    for para in original_doc.paragraphs:
        text = para.text.strip()
        if text:
            all_paragraphs.append(text)
    
    print(f"Found {len(all_paragraphs)} paragraphs")
    
    # Add document title
    title = new_doc.add_paragraph("INTERNSHIP REPORT", style='Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = new_doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Hospital Management System Development")
    run.font.size = Pt(16)
    run.font.italic = True
    
    # Add another subtitle
    subtitle2 = new_doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle2.add_run("King Abdulaziz Hospital - مستشفى الملك عبدالعزيز")
    run2.font.size = Pt(14)
    
    new_doc.add_paragraph()
    
    # Add Table of Contents
    toc_title = new_doc.add_heading("Table of Contents", level=1)
    add_toc(new_doc)
    new_doc.add_page_break()
    
    # Process each paragraph
    for i, text in enumerate(all_paragraphs):
        if not text:
            continue
        
        # Main Section Headers (SECTION 3, SECTION 4, etc.)
        if is_section_header(text):
            new_doc.add_page_break()
            new_doc.add_heading(text, level=1)
            continue
        
        # Subsection headers (3.1, 3.2, 4.1, etc.)
        if is_subsection_header(text):
            new_doc.add_heading(text, level=2)
            continue
        
        # Module headers (3.3.1, 3.3.2, etc.)
        if is_module_header(text):
            new_doc.add_heading(text, level=3)
            continue
        
        # Month headers
        if is_month_header(text):
            new_doc.add_heading(text, level=1)
            continue
        
        # Week headers
        if is_week_header(text):
            new_doc.add_heading(text, level=2)
            continue
        
        # Question headers (Q3.1.1, Q4.2.3, etc.)
        if is_question_header(text):
            new_doc.add_heading(text, level=3)
            continue
        
        # Category headers (Backend, Frontend, etc.)
        if is_category_header(text):
            para = new_doc.add_paragraph()
            run = para.add_run(text)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 102, 51)
            continue
        
        # Bullet points
        if text.startswith('-') or text.startswith('•') or text.startswith('*'):
            clean_text = text.lstrip('-•* ').strip()
            new_doc.add_paragraph(clean_text, style='List Bullet')
            continue
        
        # Numbered items
        if len(text) > 2 and text[0].isdigit() and text[1] in '.):':
            new_doc.add_paragraph(text, style='List Number')
            continue
        
        # Bold headers ending with colon
        if text.endswith(':') and len(text) < 100:
            para = new_doc.add_paragraph()
            run = para.add_run(text)
            run.bold = True
            continue
        
        # Check for inline bold patterns like "**text**"
        if '**' in text:
            para = new_doc.add_paragraph()
            parts = text.split('**')
            for j, part in enumerate(parts):
                if part:
                    run = para.add_run(part)
                    if j % 2 == 1:  # Odd parts are bold
                        run.bold = True
            continue
        
        # Regular paragraph
        new_doc.add_paragraph(text)
    
    # Save document
    print(f"Saving to: {output_path}")
    new_doc.save(output_path)
    
    print("\n" + "="*60)
    print("✅ Document reformatted successfully!")
    print("="*60)
    print(f"\nOutput file: {output_path}")
    print("\n📋 To generate the Table of Contents:")
    print("   1. Open the document in Microsoft Word")
    print("   2. Click on the Table of Contents area")
    print("   3. Press F9 or right-click → 'Update Field'")
    print("   4. Select 'Update entire table'")
    print("="*60)

if __name__ == "__main__":
    main()









