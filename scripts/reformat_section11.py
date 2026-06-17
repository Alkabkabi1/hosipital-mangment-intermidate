"""
Script to reformat SECTION 11.docx into a well-structured document
with proper headings for Table of Contents generation.
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_heading_style(doc, heading_num, font_name='Arial', font_size=14, bold=True):
    """Configure heading style"""
    style = doc.styles[f'Heading {heading_num}']
    style.font.name = font_name
    style.font.size = Pt(font_size)
    style.font.bold = bold
    return style

def add_toc(doc):
    """Add a Table of Contents field to the document"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update Table of Contents"
    
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    r_element.append(fldChar4)
    
    return paragraph

def main():
    # Paths
    input_path = r"C:\Users\sqlcc\OneDrive\Desktop\THE COPY\project-root_server_v\SECTION 11.docx"
    output_path = r"C:\Users\sqlcc\OneDrive\Desktop\THE COPY\project-root_server_v\SECTION 11 - Formatted.docx"
    
    # Read the original document
    print(f"Reading document: {input_path}")
    original_doc = Document(input_path)
    
    # Create new document
    new_doc = Document()
    
    # Configure styles
    # Title style
    title_style = new_doc.styles['Title']
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    
    # Heading styles
    h1_style = new_doc.styles['Heading 1']
    h1_style.font.name = 'Arial'
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    
    h2_style = new_doc.styles['Heading 2']
    h2_style.font.name = 'Arial'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    
    h3_style = new_doc.styles['Heading 3']
    h3_style.font.name = 'Arial'
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    
    # Normal style
    normal_style = new_doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    
    # Extract all text from original document
    all_paragraphs = []
    for para in original_doc.paragraphs:
        text = para.text.strip()
        if text:
            all_paragraphs.append(text)
    
    print(f"Found {len(all_paragraphs)} paragraphs in original document")
    
    # Patterns to identify section types
    section_patterns = {
        'main_title': ['SECTION', 'القسم'],
        'heading1': ['Q', 'Month', 'Week', 'شهر', 'أسبوع'],
        'heading2': ['Q4.', 'Q7.', 'المراجع', 'المهام'],
        'subheading': ['Backend', 'Frontend', 'Database', 'System Analysis', 'Maintenance', 'Quality'],
    }
    
    # Add document title
    title_para = new_doc.add_paragraph("SECTION 11 - INTERNSHIP REPORT", style='Title')
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = new_doc.add_paragraph("Hospital Management System Development")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    new_doc.add_paragraph()  # Empty line
    
    # Add Table of Contents placeholder
    toc_heading = new_doc.add_paragraph("Table of Contents", style='Heading 1')
    add_toc(new_doc)
    new_doc.add_page_break()
    
    # Process paragraphs and apply proper formatting
    current_section = None
    current_month = None
    current_week = None
    
    for i, text in enumerate(all_paragraphs):
        # Skip empty text
        if not text:
            continue
        
        # Detect main section headers (SECTION X)
        if text.upper().startswith('SECTION') or 'القسم' in text:
            new_doc.add_heading(text, level=1)
            current_section = text
            continue
        
        # Detect Month headers
        if text.lower().startswith('month') or 'شهر' in text:
            new_doc.add_heading(text, level=1)
            current_month = text
            continue
        
        # Detect Week headers
        if text.lower().startswith('week') or 'أسبوع' in text or 'الأسبوع' in text:
            new_doc.add_heading(text, level=2)
            current_week = text
            continue
        
        # Detect Question headers (Q4.1.1, Q7.1, etc.)
        if text.startswith('Q') and ('.' in text[:5] or ':' in text[:10]):
            # Extract question number and text
            if ':' in text:
                parts = text.split(':', 1)
                q_num = parts[0].strip()
                q_text = parts[1].strip() if len(parts) > 1 else ''
                new_doc.add_heading(f"{q_num}: {q_text}", level=3)
            else:
                new_doc.add_heading(text, level=3)
            continue
        
        # Detect category headers (Backend, Frontend, etc.)
        category_keywords = ['Backend:', 'Frontend:', 'Database:', 'System Analysis:', 
                           'Maintenance:', 'Quality Assurance:', 'Integration:',
                           'Backend tasks', 'Frontend tasks', 'Database tasks']
        is_category = False
        for keyword in category_keywords:
            if text.startswith(keyword) or text.lower().startswith(keyword.lower()):
                # Bold category with content
                para = new_doc.add_paragraph()
                run = para.add_run(text)
                run.bold = True
                is_category = True
                break
        
        if is_category:
            continue
        
        # Detect bullet points
        if text.startswith('-') or text.startswith('•') or text.startswith('*'):
            para = new_doc.add_paragraph(text[1:].strip(), style='List Bullet')
            continue
        
        # Detect numbered items
        if len(text) > 2 and text[0].isdigit() and (text[1] == '.' or text[1] == ')'):
            para = new_doc.add_paragraph(text, style='List Number')
            continue
        
        # Detect sub-headers with colons (like "Books:", "Official Documentation:")
        if ':' in text and len(text.split(':')[0]) < 50 and text.endswith(':'):
            para = new_doc.add_paragraph()
            run = para.add_run(text)
            run.bold = True
            continue
        
        # Regular paragraph
        new_doc.add_paragraph(text)
    
    # Save the new document
    print(f"Saving formatted document to: {output_path}")
    new_doc.save(output_path)
    print("Done! Document has been reformatted.")
    print("\nTo generate the Table of Contents in Word:")
    print("1. Open the document in Microsoft Word")
    print("2. Click on the Table of Contents area")
    print("3. Press F9 or right-click and select 'Update Field'")
    print("4. Choose 'Update entire table'")

if __name__ == "__main__":
    main()









