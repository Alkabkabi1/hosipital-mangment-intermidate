from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_colored_document():
    """
    Creates a Word document with colored text based on heading levels.
    Heading 1: Blue
    Heading 2: Dark Green
    Heading 3: Purple
    Heading 4: Dark Orange
    """
    
    doc = Document()
    
    # Define colors for different heading levels
    colors = {
        1: RGBColor(0, 112, 192),      # Blue
        2: RGBColor(0, 128, 0),        # Dark Green
        3: RGBColor(112, 48, 160),     # Purple
        4: RGBColor(192, 80, 77)       # Dark Orange/Red
    }
    
    def add_colored_heading(text, level=1):
        """Add a heading with colored text"""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = colors[level]
        run.font.size = Pt(16 - (level * 2))  # Decreasing size for deeper levels
        return p
    
    def add_text(text):
        """Add regular paragraph text"""
        return doc.add_paragraph(text)
    
    # Introduction
    add_colored_heading("Introduction", 1)
    add_text("The co-operative training program is a key requirement for completing the Software Engineering degree at Umm Al-Qura University, as it provides an opportunity to apply theoretical knowledge in a real work environment. My training was completed at the Information Technology (IT) Department of King Abdulaziz Hospital in Makkah, a government hospital operating under the Ministry of Health.")
    add_text("Over the training period, I participated in a variety of activities that combined full-stack web development with IT support and maintenance. I contributed to the development and enhancement of a hospital management and administrative workflow system, as well as to a maintenance and data center management system. The internship allowed me to engage in all major stages of the software development lifecycle: requirements gathering, analysis, database and system design, implementation, testing, deployment, and maintenance.")
    add_text("This report presents an overview of the training organization, followed by a detailed description of the tasks and projects I worked on, the technical and professional skills I developed, the main challenges I faced, and the lessons learned from this experience.")
    
    doc.add_paragraph()  # Spacing
    
    # About the Organization
    add_colored_heading("About the Organization", 1)
    add_text("King Abdulaziz Hospital is located in Al-Zaher, Makkah, and operates under the supervision of the Ministry of Health (MOH). It plays an important role in providing healthcare services, including outpatient clinics, inpatient care, emergency services, and specialized medical support to the local community. As part of the national healthcare network, the hospital follows MOH standards, policies, and digital transformation initiatives that align with the Kingdom's Vision 2030. The hospital has been serving patients for more than 75 years and continues to develop its medical and administrative services to meet modern healthcare requirements.")
    
    # The IT Department
    add_colored_heading("The IT Department", 2)
    add_text("The IT Department at King Abdulaziz Hospital is responsible for planning, operating, and maintaining the hospital's technical infrastructure. Its work covers a wide range of activities, including managing the data center, maintaining servers and networks, supporting clinical and administrative systems, handling user support and incident tickets, and coordinating with external vendors when needed. The department ensures that hardware, software, and communication networks operate reliably and securely so that healthcare staff can perform their duties without interruption.")
    
    # My Role in the Department
    add_colored_heading("My Role in the Department", 2)
    add_text("During my training, I rotated between different responsibilities within the IT Department, mainly in two areas:")
    
    p = doc.add_paragraph("Software development and system enhancement: ", style='List Bullet')
    p.add_run("I worked on a hospital management and administrative workflow system, as well as a maintenance and data center management system. My tasks included requirements gathering, database design, front-end and back-end development, testing, and deployment on the hospital server. I also contributed to developing and refining multiple electronic forms, approval workflows, and dashboards for different user roles.")
    
    p = doc.add_paragraph("IT support and maintenance: ", style='List Bullet')
    p.add_run("I joined IT staff on maintenance rounds in the server room and hospital departments. I assisted in checking servers, handling basic network-related issues, and maintaining printers and other devices. This exposure helped me understand how infrastructure, applications, and end users are connected in a real hospital environment.")
    
    add_text("Through these responsibilities, I gained a clear understanding of how the IT Department supports both medical and administrative operations, and how software solutions are designed and maintained in a critical healthcare context.")
    
    doc.add_paragraph()  # Spacing
    
    # Tasks, Projects, and Academic Knowledge
    add_colored_heading("Tasks, Projects, and Academic Knowledge", 1)
    add_text("This section presents the practical work completed during my co-op training at King Abdulaziz Hospital. The work is organized by month and week, and for each week I describe the tasks performed, how they relate to academic knowledge from my Software Engineering program at Umm Al-Qura University, and the skills learned.")
    
    doc.add_paragraph()
    
    # Month 1
    add_colored_heading("Month 1 – Foundations of the Hospital System", 2)
    
    # Week 1
    add_colored_heading("Week 1 – Orientation and Requirement Gathering", 3)
    
    add_colored_heading("Tasks Performed:", 4)
    add_text("During the first week, I received a general orientation about the hospital and an introduction to its different departments. I went on a field tour to understand the work environment and the existing infrastructure, which helped me see how clinical and administrative departments depend on IT services. After that, I started the requirement-gathering phase for the new system. I worked with stakeholders to collect initial system requirements, then participated in a review session with the team to analyze and clarify these requirements. By the end of the week, I prepared a preliminary requirements document, refined it based on feedback, and obtained final approval of the requirements list, which became the foundation for the system design.")
    
    add_colored_heading("Link to Academic Knowledge:", 4)
    add_text("This week strongly applied concepts from Software Engineering 1 and Systems Analysis and Design, including stakeholder identification, requirement elicitation, and requirement documentation. It also touched on professional communication skills from soft-skills courses.")
    
    add_colored_heading("Skills Learned:", 4)
    doc.add_paragraph("Communicating with non-technical stakeholders in a hospital context.", style='List Bullet')
    doc.add_paragraph("Translating business needs into structured technical requirements.", style='List Bullet')
    doc.add_paragraph("Writing and refining a requirements document for a real system.", style='List Bullet')
    
    doc.add_paragraph()
    
    # Week 2
    add_colored_heading("Week 2 – Database Design and Initial Front-End", 3)
    
    add_colored_heading("Tasks Performed:", 4)
    add_text("In the second week, I focused on designing the database. I started by preparing an initial ERD (Entity–Relationship Diagram) that modeled the main entities such as users, departments, and different request types. I then reviewed this design with the team and stakeholder, received comments, and adjusted the ERD until the final design was approved.")
    add_text("After the ERD was finalized, I moved to the initial front-end work. I linked the initial pages as a front-end, which included the login page, new account (registration) page, and the personal page/homepage. I then tested the buttons and navigation flows to ensure that each button worked correctly and that users could move between pages without issues. At the end of the week, I reviewed the progress with stakeholders to gather feedback on both the design and the initial interface.")
    add_text("(Figure 1: Initial ERD for the Hospital System designed during Week 2)")
    add_text("(Figure 2: Initial login, registration, and personal homepage interfaces)")
    
    add_colored_heading("Link to Academic Knowledge:", 4)
    add_text("This week applied Database Systems (ER modeling, relationships, normalization basics) and Web Development (HTML/CSS/JavaScript for basic pages and navigation). Testing the buttons and flows connected to Software Quality and Testing principles like verifying functionality and user flows.")
    
    add_colored_heading("Skills Learned:", 4)
    doc.add_paragraph("Designing an ERD that reflects real organizational workflows.", style='List Bullet')
    doc.add_paragraph("Implementing simple user interfaces and linking pages in a coherent flow.", style='List Bullet')
    doc.add_paragraph("Performing early functional testing and incorporating stakeholder feedback.", style='List Bullet')
    
    doc.add_paragraph()
    
    # Week 3
    add_colored_heading("Week 3 – Dashboards and Request Pages", 3)
    
    add_colored_heading("Tasks Performed:", 4)
    add_text("During the third week, I worked on role-based interfaces. I added a dashboard page for the manager and another dashboard page for the employee, each showing information relevant to that role. I then created a clearance request page (إخلاء طرف) and a work resumption request page (مباشرة عمل). All new pages were linked to the employee's homepage, so employees could easily navigate between their dashboard and request pages. At the end, I tested all added pages and buttons to ensure that routing, submission, and navigation were functioning as expected.")
    add_text("(Figure 3: Manager and employee dashboard interfaces)")
    add_text("(Figure 4: Clearance request and work resumption request pages)")
    
    add_colored_heading("Link to Academic Knowledge:", 4)
    add_text("This week related to Human–Computer Interaction (HCI) and UX concepts such as clarity, consistency, and role-based design. It also built on client–server architecture ideas from software engineering courses, as I prepared front-end structures that would later connect to back-end logic.")
    
    add_colored_heading("Skills Learned:", 4)
    doc.add_paragraph("Designing dashboards tailored for different roles (manager vs employee).", style='List Bullet')
    doc.add_paragraph("Structuring request pages for common HR/administrative workflows.", style='List Bullet')
    doc.add_paragraph("Strengthening testing practice for multi-page web flows.", style='List Bullet')
    
    doc.add_paragraph()
    
    # Week 4
    add_colored_heading("Week 4 – Profiles, Inboxes, and Local Storage", 3)
    
    add_colored_heading("Tasks Performed:", 4)
    add_text("In the fourth week, I enhanced the system with more detailed user features. I added a profile page for the manager and a profile page for the employee, allowing each user type to view and manage basic information. I then added two inbox sections for the manager to view different categories of received requests, and two inbox sections for employees to view the requests they created themselves.")
    add_text("To improve responsiveness, I linked all pages with Local Storage, which was used to temporarily store data on the client side and improve user experience in certain flows. Finally, I conducted a comprehensive review with stakeholders, walking them through the system and ensuring their satisfaction before moving into deeper development.")
    add_text("(Figure 5: Profile and inbox pages for managers and employees)")
    
    add_colored_heading("Link to Academic Knowledge:", 4)
    add_text("This week applied knowledge from Web Technologies (Local Storage and client-side state), UX Design (organizing inboxes and profiles), and Software Engineering (iterative development and review cycles).")
    
    add_colored_heading("Skills Learned:", 4)
    doc.add_paragraph("Using Local Storage to support smoother interactions.", style='List Bullet')
    doc.add_paragraph("Designing intuitive inbox structures for different user types.", style='List Bullet')
    doc.add_paragraph("Presenting progress in a structured way to stakeholders and incorporating their feedback.", style='List Bullet')
    
    doc.add_paragraph()
    
    # Month 2
    add_colored_heading("Month 2 – Maintenance Exposure and System Expansion", 2)
    
    # Week 1
    add_colored_heading("Week 1 – Server Room and Maintenance Rounds", 3)
    
    add_colored_heading("Tasks Performed:", 4)
    add_text("At the start of the second month, I shifted temporarily towards infrastructure and support to better understand the environment our systems run in. I joined an orientation tour in the server room to learn about the main devices used by the hospital. I received an explanation of the internal network components and how they function, including switches, routers, and connections between departments.")
    add_text("Throughout the week, I assisted in morning maintenance rounds, which included checking servers, monitoring their status, and noting issues. I also supported printer maintenance, ensuring printers were correctly configured and connected to the network. When incident reports related to network changes came in, I helped handle them by checking connectivity and configuration. At the end of the week, I prepared a weekly report summarizing maintenance activities and incidents handled during the rounds.")
    
    add_colored_heading("Link to Academic Knowledge:", 4)
    add_text("This week built on Computer Networks (topologies, LAN, devices) and Operating Systems / System Administration concepts such as server checks and resource availability. The reporting aspect linked to IT Service Management and documentation skills.")
    
    add_colored_heading("Skills Learned:", 4)
    doc.add_paragraph("Understanding how IT infrastructure supports hospital systems.", style='List Bullet')
    doc.add_paragraph("Handling basic network- and printer-related incidents.", style='List Bullet')
    doc.add_paragraph("Writing concise maintenance and incident summary reports.", style='List Bullet')
    
    doc.add_paragraph()
    
    # Week 2
    add_colored_heading("Week 2 – Delegates Management and Department Heads", 3)
    
    add_colored_heading("Tasks Performed:", 4)
    add_text("In the second week of Month 2, I returned to system development, focusing on delegation features. I added a new page for managing delegates and linked it to the admin (manager) homepage, enabling the manager to view and manage delegation relationships. After implementing the page, I reviewed it and verified its functionality.")
    add_text("I then updated the employee page to allow employees to receive and view delegation requests. A new page was also added for assigning department heads, where their roles could be defined, and I implemented a dedicated button for the admin to manage this. To support these features, I prepared APIs to connect the delegates page with the back-end logic. Finally, I tested the system, applied necessary adjustments, and reviewed the changes with stakeholders to ensure that delegation and department-head management matched their expectations.")
    add_text("(Figure 6: Delegates management page, department heads assignment interface)")
    
    add_colored_heading("Link to Academic Knowledge:", 4)
    add_text("This week drew on Role-Based Access Control (RBAC) concepts from information security courses and API design from web development. It also connected to database design for modeling delegates, roles, and department heads.")
    
    add_colored_heading("Skills Learned:", 4)
    doc.add_paragraph("Designing and implementing delegation workflows.", style='List Bullet')
    doc.add_paragraph("Integrating front-end pages with back-end APIs.", style='List Bullet')
    doc.add_paragraph("Coordinating with stakeholders on sensitive, permission-related features.", style='List Bullet')
    
    doc.add_paragraph()
    
    # Week 3
    add_colored_heading("Week 3 – Role Assignment and Weekly Reporting", 3)
    
    add_colored_heading("Tasks Performed:", 4)
    add_text("The third week focused on role assignment. I updated the employee page to handle role assignment requests, ensuring employees could initiate or view role-related actions. I then linked the role assignment page to the employee's front-end, integrating it into the existing navigation.")
    add_text("To complete the technical side, I prepared APIs to connect the role assignment page to the back-end, so that role changes would be correctly processed and saved. After these integrations, I tested the new page thoroughly, confirming that all buttons and flows worked correctly. I then reviewed the updates with stakeholders, obtained their approval, and prepared a detailed weekly report summarizing testing results and feedback.")
    add_text("(Figure 7: Role assignment interface and employee-side view)")
    
    add_colored_heading("Link to Academic Knowledge:", 4)
    add_text("This week again related to RBAC and authorization, many-to-many database relationships (users–roles–permissions), and software testing for new functionality.")
    
    add_colored_heading("Skills Learned:", 4)
    doc.add_paragraph("Implementing role assignment logic from both the front-end and back-end perspectives.", style='List Bullet')
    doc.add_paragraph("Documenting testing outcomes and stakeholder feedback clearly.", style='List Bullet')
    doc.add_paragraph("Ensuring UI behavior dynamically reflects user roles and permissions.", style='List Bullet')
    
    doc.add_paragraph()
    
    # Week 4
    add_colored_heading("Week 4 – Database Migration and Deployment", 3)
    
    add_colored_heading("Tasks Performed:", 4)
    add_text("In the fourth week, I focused on moving from prototypes to a more production-ready system. I began by reviewing earlier project components and preparing a migration plan to shift them to the real database. I then migrated the dashboard pages so they would use the database instead of mock/local data, followed by testing the dashboard after migration to verify that data displayed correctly.")
    add_text("Next, I migrated the clearance and onboarding request pages to use the database as well. To prepare for deployment, I adjusted server settings so the system could be hosted on the hospital's server. Finally, I deployed the system on the hospital server and conducted a final test to ensure that the system functioned correctly in its new environment.")
    add_text("(Figure 8: Deployed dashboard and request pages connected to the live database)")
    
    add_colored_heading("Link to Academic Knowledge:", 4)
    add_text("This week applied database migration, deployment, and environment configuration concepts related to DevOps and system administration. It also reinforced regression testing ideas from software testing courses.")
    
    add_colored_heading("Skills Learned:", 4)
    doc.add_paragraph("Planning and executing database migrations.", style='List Bullet')
    doc.add_paragraph("Deploying a web system on a real hospital server.", style='List Bullet')
    doc.add_paragraph("Performing final checks and resolving environment-related issues.", style='List Bullet')
    
    doc.add_paragraph()
    
    # Month 3
    add_colored_heading("Month 3 – Advanced System Features and Maintenance", 2)
    
    # Week 1
    add_colored_heading("Week 1 – External Access and Core Interfaces", 3)
    
    add_colored_heading("Tasks Performed:", 4)
    add_text("The first week of Month 3 focused on building external entry points and improving the basic user experience. I implemented Employee Login & Authentication, ensuring that employees could securely log into the system. I also developed an External Request Form for users outside the internal network to submit maintenance or service requests. To allow users to follow up, I created a Status Tracking Page, where users could track the progress of their requests.")
    add_text("In addition, I built an Exit Request Form to handle work termination or exit procedures. I worked on Arabic/English translation within the interface to support bilingual users, and I improved the responsive UI design so the system would adapt better to different screen sizes and devices.")
    add_text("(Figure 9: Employee login, external request, and status tracking interfaces)")
    
    add_colored_heading("Link to Academic Knowledge:", 4)
    add_text("This week applied web authentication concepts, form design, and responsiveness from web development courses. It also touched on localization/internationalization topics for supporting multiple languages.")
    
    add_colored_heading("Skills Learned:", 4)
    doc.add_paragraph("Implementing secure login and basic authentication flows.", style='List Bullet')
    doc.add_paragraph("Designing request and status tracking interfaces that are user friendly.", style='List Bullet')
    doc.add_paragraph("Creating bilingual, responsive web pages.", style='List Bullet')
    
    doc.add_paragraph()
    
    # Week 2
    add_colored_heading("Week 2 – Admin Tools, Internal Tickets, and Reports", 3)
    
    add_colored_heading("Tasks Performed:", 4)
    add_text("In the second week, I moved to admin-facing features. I developed an Admin Dashboard with RBAC (role-based access control), allowing admins to see high-level information and access modules based on their roles. I then created a Request Review Dashboard for admins to review, approve, or reject requests.")
    add_text("Additionally, I implemented an Internal Ticket System, allowing internal staff to log issues and track them separately from external requests. To support monitoring and decision making, I developed Statistics & Reports pages that aggregate and visualize key information, and I built functionality for Admin Exit Management, enabling admins to manage exit-related processes for employees. I also worked on API Backend Integration to ensure the admin views were powered by real data from the server.")
    add_text("(Figure 10: Admin RBAC dashboard, request review, and statistics views)")
    
    add_colored_heading("Link to Academic Knowledge:", 4)
    add_text("This week leveraged RBAC and security, data visualization and reporting, and RESTful API integration concepts from software engineering and web courses.")
    
    add_colored_heading("Skills Learned:", 4)
    doc.add_paragraph("Designing and implementing dashboards for administrators.", style='List Bullet')
    doc.add_paragraph("Integrating front-end components tightly with API back-ends.", style='List Bullet')
    doc.add_paragraph("Managing internal vs external ticket flows in one system.", style='List Bullet')
    
    doc.add_paragraph()
    
    # Week 3
    add_colored_heading("Week 3 – Data Center and Server Management Features", 3)
    
    add_colored_heading("Tasks Performed:", 4)
    add_text("The third week of Month 3 focused on the data center. I created a Data Center Dashboard that provides an overview of key infrastructure elements. I then developed a Server Inventory Page that lists servers and their properties. To support operations, I added Server Edit/Delete Features, allowing admins to update or remove records as needed.")
    add_text("I also integrated Server Monitoring Metrics to display basic health information, and extended monitoring to AC System Monitoring to track cooling conditions. On top of that, I built an Alert Management System to notify relevant staff when certain thresholds or error conditions are reached.")
    add_text("(Figure 11: Data center dashboard and server inventory/monitoring screens)")
    
    add_colored_heading("Link to Academic Knowledge:", 4)
    add_text("This week brought together concepts from Operating Systems, Computer Networks, and system monitoring. It also used CRUD operations and data validation techniques from database and web programming courses.")
    
    add_colored_heading("Skills Learned:", 4)
    doc.add_paragraph("Modeling technical assets (servers, AC units) in a database.", style='List Bullet')
    doc.add_paragraph("Designing inventory and monitoring interfaces for IT staff.", style='List Bullet')
    doc.add_paragraph("Implementing alert mechanisms to support proactive incident handling.", style='List Bullet')
    
    doc.add_paragraph()
    
    # Week 4
    add_colored_heading("Week 4 – Backup, Power, Scheduling, and Incident Handling", 3)
    
    add_colored_heading("Tasks Performed:", 4)
    add_text("In the final week of Month 3, I worked on reliability and operations. I implemented Backup Status Display pages to show the current backup state. I then set up Backup Scheduling so backups could be run periodically. I added Power & UPS Management interfaces to monitor power conditions and UPS status.")
    add_text("To support routine work, I created a Maintenance Scheduler to schedule recurring maintenance tasks. I also built an Incident Report System that allows staff to log incidents, categorize them, and track resolutions. Finally, I performed Final Testing & Debugging of all these modules to ensure they worked together without conflicts.")
    add_text("(Figure 12: Backup, power/UPS, scheduler, and incident-reporting interfaces)")
    
    add_colored_heading("Link to Academic Knowledge:", 4)
    add_text("This week reflected topics from system reliability, fault tolerance, scheduling, and incident management. It also applied software testing methods to validate integrated modules.")
    
    add_colored_heading("Skills Learned:", 4)
    doc.add_paragraph("Designing features that improve system reliability and resilience.", style='List Bullet')
    doc.add_paragraph("Scheduling automated or routine tasks in an operational environment.", style='List Bullet')
    doc.add_paragraph("Debugging interconnected modules and verifying system behavior end-to-end.", style='List Bullet')
    
    doc.add_paragraph()
    
    # Month 4
    add_colored_heading("Month 4 – Enhancements, Optimization, and Final Feature Set", 2)
    
    # Week 1
    add_colored_heading("Week 1 – Backend Structure and API Refinement", 3)
    
    add_colored_heading("Tasks Performed:", 4)
    add_text("During the first week of Month 4, I focused on improving the internal structure of the system. I enhanced the backend structure and refined API routing to improve module performance and clarity. I reviewed existing system components, making code quality adjustments and cleaning up duplicated or unclear logic. I also optimized deployment configuration, preparing the server environment for smoother updates and potential expansions. Finally, I enhanced workflow logic in preparation for upcoming features and conducted general system checks to ensure stability and integration readiness.")
    
    add_colored_heading("Link to Academic Knowledge:", 4)
    add_text("This week applied refactoring, modularity, and API design concepts from advanced software engineering courses, as well as basic DevOps ideas related to deployment and configuration.")
    
    add_colored_heading("Skills Learned:", 4)
    doc.add_paragraph("Improving performance and maintainability through code refactoring.", style='List Bullet')
    doc.add_paragraph("Managing APIs in a structured, predictable way.", style='List Bullet')
    doc.add_paragraph("Preparing production environments for feature growth.", style='List Bullet')
    
    doc.add_paragraph()
    
    # Week 2
    add_colored_heading("Week 2 – Workflow Enhancements and UI/UX Improvements", 3)
    
    add_colored_heading("Tasks Performed:", 4)
    add_text("In the second week, I improved employee request workflows and expanded form functionalities based on feedback. I refactored request modules to support additional request types, making them easier to extend. I also enhanced notifications and commissioner workflow logic to better match actual approval processes. Furthermore, I improved job description and credential management interfaces, making them clearer and easier to use. I updated system UI components to improve usability and visual consistency across pages.")
    add_text("(Figure 13: Updated forms, notifications, and job-description/credential interfaces)")
    
    add_colored_heading("Link to Academic Knowledge:", 4)
    add_text("This week aligned with UX/UI principles, workflow modeling, and maintainable software design. It also applied iterative improvement ideas from Agile methodologies.")
    
    add_colored_heading("Skills Learned:", 4)
    doc.add_paragraph("Evolving existing features based on real feedback.", style='List Bullet')
    doc.add_paragraph("Designing more flexible form and workflow structures.", style='List Bullet')
    doc.add_paragraph("Improving user experience through visual and interaction consistency.", style='List Bullet')
    
    doc.add_paragraph()
    
    # Week 3
    add_colored_heading("Week 3 – Multi-Approval Workflows and Security Improvements", 3)
    
    add_colored_heading("Tasks Performed:", 4)
    add_text("The third week involved deeper logic and quality assurance. I enhanced multi-approval workflows to support more complex approval chains, such as sequential approvals and different roles participating. I tested unified inbox systems and improved how the admin dashboard behaves with various scenarios. I then ran 67 automated tests and verified that they all passed, proving the stability of core logic. I also performed security improvements, including Content Security Policy (CSP), rate limiting, and tracing to monitor potentially suspicious activity. Finally, I reviewed the database structure and migration scripts to optimize schema relations and ensure consistency.")
    
    add_colored_heading("Link to Academic Knowledge:", 4)
    add_text("This week strongly connected to software testing (unit, integration), information security, and database design optimization.")
    
    add_colored_heading("Skills Learned:", 4)
    doc.add_paragraph("Implementing and validating complex multi-approval workflows.", style='List Bullet')
    doc.add_paragraph("Writing and interpreting automated test results.", style='List Bullet')
    doc.add_paragraph("Applying practical security measures to a real web system.", style='List Bullet')
    
    doc.add_paragraph()
    
    # Week 4
    add_colored_heading("Week 4 – Final Enhancements and Documentation", 3)
    
    add_colored_heading("Tasks Performed:", 4)
    add_text("In the final week of Month 4, I completed enhancements for all employee request types (17+ forms in total), ensuring that each one worked end-to-end. I improved commissioner ticketing workflows and their approval routes. I finalized enhancements to job description and credentials modules to ensure accurate and reliable usage. I also created and completed the Project Management Plan (PMP) document, describing scope, timelines, risks, and roles. Then I prepared the User Manual with full workflow instructions and screenshots, enabling non-technical users to understand the system. Lastly, I conducted a final review of all enhancements and prepared the system for delivery.")
    
    add_colored_heading("Link to Academic Knowledge:", 4)
    add_text("This week integrated project management concepts (PMP structure, scope and risk management) and technical documentation skills. It also brought together everything from requirements, design, implementation, testing, and deployment into a final deliverable.")
    
    add_colored_heading("Skills Learned:", 4)
    doc.add_paragraph("Completing and stabilizing a large set of interconnected forms and workflows.", style='List Bullet')
    doc.add_paragraph("Writing structured project documentation and user manuals.", style='List Bullet')
    doc.add_paragraph("Preparing a system for handover and potential production use.", style='List Bullet')
    
    doc.add_paragraph()
    
    # Extra Week
    add_colored_heading("Extra Week – Final Polishing and Job Description Feature", 2)
    
    add_colored_heading("Tasks Performed:", 4)
    add_text("During the extra training week, I focused on final polishing and one key feature. I revisited system flows to fix minor issues and improve clarity in navigation and messages. I also gave special attention to the Job Description feature, where an employee can submit their job description, and the admin reviews and approves it. Once approved, the job description appears to the employee who submitted it, providing a clear, traceable record.")
    add_text("Additionally, during this period I finalized and refined the PMP Word document and the User Manual Word document, ensuring they accurately reflected the final state of the system and included all the major form categories:")
    
    p = doc.add_paragraph("Core Workflow Forms (3): ", style='List Bullet')
    p.add_run("Clearance (إخلاء طرف), Onboarding (مباشرة عمل), Delegation (تفويض).")
    
    p = doc.add_paragraph("Certificate Requests (2): ", style='List Bullet')
    p.add_run("Certificate of Identification (شهادة تعريف), Experience Certificate (شهادة خبرة).")
    
    p = doc.add_paragraph("Assignment & Transfer Forms (3): ", style='List Bullet')
    p.add_run("Assignment (قرار تكليف), Assignment Termination (إنهاء تكليف), Internal Transfer (نقل داخلي).")
    
    p = doc.add_paragraph("Leave Requests (2): ", style='List Bullet')
    p.add_run("Leave Request (إجازة), Maternity Leave (طلب إجازة رعاية مولود).")
    
    p = doc.add_paragraph("Housing & Allowance Forms (2): ", style='List Bullet')
    p.add_run("Housing Allowance for Saudi Doctors (بدل سكن أطباء سعوديين), Contractor Housing (بدل سكن المتعاقدين).")
    
    p = doc.add_paragraph("Travel & Transportation Forms (4): ", style='List Bullet')
    p.add_run("Travel Order for non-Saudis, Airlines Ticket Request, Saudi Ticket Compensation, Ticket Compensation for Contractors.")
    
    p = doc.add_paragraph("Guarantee Forms (3): ", style='List Bullet')
    p.add_run("Detailed Guarantee (كفالة غرم وأداء وحضور بديل), Guarantee Fine & Performance (كفالة غرم وأداء), Public Law Guarantee (كفالة غرم وأداء في الحق العام).")
    
    p = doc.add_paragraph("Financial & Compensation Forms (1): ", style='List Bullet')
    p.add_run("Reward/Entitlement Refund (صرف مستحقات).")
    
    p = doc.add_paragraph("Exit & Administrative Forms (2): ", style='List Bullet')
    p.add_run("Exit Request (طلب إنهاء العمل), Job Description (الوصف الوظيفي).")
    
    add_text("(Figure 14: Example interface for job description submission and admin approval)")
    
    add_colored_heading("Link to Academic Knowledge:", 4)
    add_text("This final period tied back to requirements traceability, HR and administrative workflows, project management, and technical communication. The job-description approval loop reflected real-world workflow and authorization modeling.")
    
    add_colored_heading("Skills Learned:", 4)
    doc.add_paragraph("Finalizing a complex system and aligning it with documentation and real use cases.", style='List Bullet')
    doc.add_paragraph("Designing and implementing a complete submission–review–approval cycle.", style='List Bullet')
    doc.add_paragraph("Organizing and categorizing a large number of forms in a clear, maintainable structure.", style='List Bullet')
    
    doc.add_paragraph()
    
    # Problems and Challenges
    add_colored_heading("Problems and Challenges", 1)
    add_text("During the training period, I faced several challenges that contributed to my learning and personal growth. These challenges can be grouped into three main categories: technical, time-related, and organizational/administrative.")
    
    add_colored_heading("Challenges Related to New Technologies", 2)
    add_text("At the beginning of the internship, I had not previously built a complete full-stack application on my own. Working on a real hospital system required me to connect many concepts from different university courses—such as databases, web development, software engineering, and networks—and apply them together in one project. Understanding and implementing multi-step workflows, role-based access, and deployment on a real server were initially challenging.")
    add_text("To overcome these difficulties, I relied on official documentation, online resources, and continuous experimentation. I also benefited from the guidance of my trainer and the IT team, and I gradually became more confident in designing and implementing complex features. Each issue I faced, whether in database design, front-end behavior, or back-end integration, became an opportunity to deepen my technical understanding.")
    
    add_colored_heading("Time Management and Task Coordination", 2)
    add_text("At the start of the training, I was involved in both maintenance rounds and software development tasks at the same time. Switching between these two areas made it difficult to fully focus on either, and sometimes slowed down progress in development. Later, the team was reorganized so that some trainees focused mainly on development while others focused on maintenance. This change helped improve time management and allowed more consistent progress on the system.")
    add_text("Another challenge was related to changing or unclear requirements. In some cases, stakeholders were not fully sure about what they wanted at the beginning, or they requested additional changes after features had already been implemented. This sometimes led to rework and extra effort, especially when changes affected core workflows or required learning new technologies. Over time, I learned to ask more detailed questions, document requirements clearly, and confirm them before starting implementation.")
    
    add_colored_heading("Challenges Related to Training Structure and Procedures", 2)
    add_text("At the beginning of the training, the overall expectations were not completely clear. It was not always obvious what the exact priorities were or which tasks should be done first, which caused some confusion in the early weeks. With time, and through regular communication with the trainer, the tasks and goals became clearer.")
    add_text("Towards the end of the training, there was also some uncertainty about the completion procedures—for example, whether the training ends strictly after finishing the required hours or at the official end date, and which documents were needed from the trainer. This ambiguity led to continuing work even after completing the required hours. Although this was a challenge, it also taught me the importance of asking precise administrative questions early and confirming formal requirements with the university and the training organization.")
    add_text("Despite these difficulties, each challenge contributed to developing my problem-solving abilities, communication skills, and adaptability in a real work environment.")
    
    doc.add_paragraph()
    
    # Conclusion
    add_colored_heading("Conclusion", 1)
    add_text("My internship at the IT Department of King Abdulaziz Hospital was a valuable and transformative experience that helped bridge the gap between theoretical study and practical application. Through my work on the hospital management and maintenance systems, I was able to participate in the full software development lifecycle, from requirements gathering and system design to implementation, testing, deployment, and maintenance. I also gained direct exposure to IT operations in a hospital environment, including server room activities, network-related issues, and device maintenance.")
    add_text("This experience significantly improved my technical skills in full-stack development, database design, API integration, and system deployment, as well as my ability to read documentation, debug issues, and adopt best practices. At the same time, I developed important professional skills such as time management, teamwork, communication with stakeholders, and technical documentation. Working on real systems that support daily hospital operations gave me a clearer understanding of the responsibilities of software engineers in critical sectors like healthcare.")
    add_text("Overall, the internship strengthened my confidence in pursuing a career in software development and system design. It showed me how the concepts learned at university can be turned into practical solutions that solve real problems and support real users. I am grateful for the opportunity to train at King Abdulaziz Hospital and for the guidance provided by the IT team and my trainer, and I consider this experience an important step in my academic and professional journey.")
    
    doc.add_paragraph()
    
    # References
    add_colored_heading("References", 1)
    add_text("You can use or adapt these as references in your report:")
    
    doc.add_paragraph("Ministry of Health (Saudi Arabia) – official website and digital transformation materials.", style='List Bullet')
    doc.add_paragraph("King Abdulaziz Hospital (Makkah) – internal training documents and system requirements provided by the IT Department.", style='List Bullet')
    doc.add_paragraph("Pressman, R. S., & Maxim, B. R. Software Engineering: A Practitioner's Approach.", style='List Bullet')
    doc.add_paragraph("Sommerville, I. Software Engineering (10th Edition).", style='List Bullet')
    doc.add_paragraph("MySQL – Official Documentation (dev.mysql.com).", style='List Bullet')
    doc.add_paragraph("MDN Web Docs – HTML, CSS, and JavaScript references (developer.mozilla.org).", style='List Bullet')
    doc.add_paragraph("Node.js – Official Documentation (nodejs.org).", style='List Bullet')
    doc.add_paragraph("Express.js – Official Documentation (expressjs.com).", style='List Bullet')
    
    add_text("If your university wants Arabic references as well, you can also add:")
    doc.add_paragraph("وزارة الصحة – بوابة وزارة الصحة السعودية والمبادرات الرقمية.", style='List Bullet')
    
    doc.add_paragraph()
    
    # Appendix
    add_colored_heading("Appendix (Structure Suggestion)", 1)
    add_text("You can structure the appendix like this:")
    
    add_colored_heading("Appendix A — System Screenshots", 2)
    add_text("Includes screenshots of key system pages such as:")
    
    doc.add_paragraph("Login and registration screens", style='List Bullet')
    doc.add_paragraph("Employee and manager dashboards", style='List Bullet')
    doc.add_paragraph("Clearance, onboarding, and other request forms", style='List Bullet')
    doc.add_paragraph("Delegation and department head management pages", style='List Bullet')
    doc.add_paragraph("Role assignment interfaces", style='List Bullet')
    doc.add_paragraph("Server inventory and data center dashboard", style='List Bullet')
    doc.add_paragraph("AC monitoring and alert system", style='List Bullet')
    doc.add_paragraph("Backup and UPS management pages", style='List Bullet')
    doc.add_paragraph("Maintenance scheduler", style='List Bullet')
    doc.add_paragraph("Incident report system", style='List Bullet')
    doc.add_paragraph("Job description submission and approval screens", style='List Bullet')
    
    add_colored_heading("Appendix B — GitHub Repository", 2)
    doc.add_paragraph("Link to the project's GitHub repository containing:", style='List Bullet')
    p = doc.add_paragraph("Source code", style='List Bullet 2')
    p = doc.add_paragraph("Configuration files", style='List Bullet 2')
    p = doc.add_paragraph("ERD diagrams", style='List Bullet 2')
    p = doc.add_paragraph("Selected documentation", style='List Bullet 2')
    
    add_colored_heading("Appendix C — Training Plan", 2)
    doc.add_paragraph("Official training plan", style='List Bullet')
    doc.add_paragraph("Weekly student report forms", style='List Bullet')
    doc.add_paragraph("Any documents related to the structure and duration of the co-op training", style='List Bullet')
    
    add_colored_heading("Appendix D — Additional Documents", 2)
    doc.add_paragraph("Project Management Plan (PMP) – only main pages or table of contents", style='List Bullet')
    doc.add_paragraph("User Manual – table of contents or selected screenshots", style='List Bullet')
    doc.add_paragraph("Any additional architecture or workflow diagrams", style='List Bullet')
    
    # Save the document
    doc.save('Training_Report_Colored.docx')
    print("✅ Document created successfully: Training_Report_Colored.docx")
    print("\nColor scheme used:")
    print("  - Heading 1 (Main sections): Blue")
    print("  - Heading 2 (Subsections): Dark Green")
    print("  - Heading 3 (Weeks): Purple")
    print("  - Heading 4 (Task categories): Dark Orange/Red")

if __name__ == "__main__":
    create_colored_document()

