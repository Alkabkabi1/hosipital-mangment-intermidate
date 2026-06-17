"""
Quick Diagram Export - Uses online Mermaid API (NO CLI needed!)
This is the EASIEST method - just run this script!

Requirements: pip install python-docx requests
"""

import os
import re
import base64
import zlib
import requests
from docx import Document
from docx.shared import Inches
from io import BytesIO

def extract_mermaid_diagrams(md_file):
    """Extract all mermaid diagrams from markdown file"""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    pattern = r'## (Diagram \d+: [^\n]+)\n\n```mermaid\n(.*?)```'
    matches = re.findall(pattern, content, re.DOTALL)
    
    diagrams = []
    for title, code in matches:
        diagrams.append({
            'title': title,
            'code': code.strip()
        })
    
    return diagrams

def encode_mermaid_for_url(mermaid_code):
    """Encode mermaid code for mermaid.ink URL"""
    # Convert to bytes
    json_string = '{"code":"' + mermaid_code.replace('\n', '\\n').replace('"', '\\"') + '","mermaid":{"theme":"default"}}'
    
    # Compress and encode
    compressed = zlib.compress(json_string.encode('utf-8'))
    encoded = base64.urlsafe_b64encode(compressed).decode('ascii')
    
    return encoded

def render_diagram_online(mermaid_code):
    """Render diagram using mermaid.ink online service"""
    try:
        # Method 1: Using mermaid.ink
        encoded = encode_mermaid_for_url(mermaid_code)
        url = f"https://mermaid.ink/img/{encoded}?type=png&width=1200"
        
        response = requests.get(url, timeout=30)
        
        if response.status_code == 200:
            return BytesIO(response.content)
        else:
            print(f"  Status code: {response.status_code}")
            return None
            
    except Exception as e:
        print(f"  Error: {e}")
        return None

def insert_diagrams_into_word(doc_path, diagrams, output_path):
    """Insert rendered diagrams into Word document"""
    print(f"\nProcessing: {doc_path}")
    
    try:
        doc = Document(doc_path)
    except Exception as e:
        print(f"  X Error opening document: {e}")
        return False
    
    # Find appendix or add at end
    doc.add_page_break()
    doc.add_paragraph('Appendix A: System Architecture Diagrams', style='Heading 1')
    doc.add_paragraph('The following diagrams illustrate the system architecture, workflows, and technical design.')
    doc.add_paragraph()
    
    # Process each diagram
    success_count = 0
    for i, diagram in enumerate(diagrams, 1):
        print(f"  [{i}/{len(diagrams)}] {diagram['title']}...")
        
        # Add title
        doc.add_paragraph(diagram['title'], style='Heading 2')
        
        # Render and insert image
        image_data = render_diagram_online(diagram['code'])
        
        if image_data:
            try:
                doc.add_picture(image_data, width=Inches(6.5))
                doc.add_paragraph()  # Spacing
                print(f"      > Inserted successfully")
                success_count += 1
            except Exception as e:
                print(f"      X Failed to insert: {e}")
        else:
            print(f"      X Failed to render")
            doc.add_paragraph(f"[Diagram could not be rendered - please insert manually]")
    
    # Save
    try:
        doc.save(output_path)
        print(f"  > Saved: {output_path} ({success_count}/{len(diagrams)} diagrams inserted)")
        return True
    except Exception as e:
        print(f"  X Error saving: {e}")
        return False

def main():
    print("=" * 70)
    print("QUICK DIAGRAM EXPORT - Using Online API")
    print("=" * 70)
    print("\nThis will:")
    print("  1. Extract all 8 diagrams from DIAGRAMS_MERMAID.md")
    print("  2. Render them using mermaid.ink online service")
    print("  3. Insert them into your Word documents")
    print("\nNo installation needed! Just Python + internet connection.")
    print("=" * 70)
    
    # Extract diagrams
    print("\n[STEP 1] Extracting diagrams...")
    try:
        diagrams = extract_mermaid_diagrams('DIAGRAMS_MERMAID.md')
        print(f"   > Found {len(diagrams)} diagrams")
    except FileNotFoundError:
        print("   X ERROR: DIAGRAMS_MERMAID.md not found!")
        print("   Make sure you're running this from the project directory.")
        return
    except Exception as e:
        print(f"   X Error: {e}")
        return
    
    # Process documents
    print("\n[STEP 2] Processing Word documents...")
    
    documents = [
        ('Hospital_Management_System_PMP_with_Budget.docx', 
         'Hospital_Management_System_PMP_with_Budget_FINAL.docx',
         'Arabic Version'),
        ('Hospital_Management_System_PMP_English.docx',
         'Hospital_Management_System_PMP_English_FINAL.docx',
         'English Version')
    ]
    
    processed = 0
    for input_doc, output_doc, description in documents:
        if os.path.exists(input_doc):
            print(f"\n{description}:")
            if insert_diagrams_into_word(input_doc, diagrams, output_doc):
                processed += 1
        else:
            print(f"\n{description}: X Not found ({input_doc})")
    
    # Summary
    print("\n" + "=" * 70)
    if processed > 0:
        print("SUCCESS!")
        print("=" * 70)
        print(f"\n{processed} document(s) created with diagrams:")
        for _, output_doc, desc in documents:
            if os.path.exists(output_doc):
                print(f"  - {output_doc}")
        print("\nYou can now use these documents!")
        print("\nTIP: Open the documents and add a Table of Contents:")
        print("   Word -> References -> Table of Contents -> Choose a style")
    else:
        print("NO DOCUMENTS PROCESSED")
        print("=" * 70)
        print("\nPlease check:")
        print("  1. Word documents exist in current directory")
        print("  2. You have internet connection (for rendering)")
        print("  3. python-docx is installed: pip install python-docx requests")
    
    print()

if __name__ == '__main__':
    try:
        import requests
        from docx import Document
        main()
    except ImportError as e:
        print("=" * 70)
        print("WARNING: MISSING REQUIRED PACKAGE")
        print("=" * 70)
        print(f"\nError: {e}")
        print("\nPlease install required packages:")
        print("  pip install python-docx requests")
        print("\nThen run this script again.")

