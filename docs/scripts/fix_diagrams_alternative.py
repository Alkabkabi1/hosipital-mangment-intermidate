"""
Alternative Diagram Export - Using Kroki API (more reliable)
Kroki is a free service that renders Mermaid diagrams
"""

import os
import re
import requests
from docx import Document
from docx.shared import Inches
from io import BytesIO

def extract_mermaid_diagrams(md_file):
    """Extract all mermaid diagrams from markdown file"""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    pattern = r'## (Diagram \d+: [^\n]+)\n\n```mermaid\n(.*?)```'
    matches = re.findall(pattern, content, re.DOTALL)
    
    diagrams = []
    for title, code in matches:
        diagrams.append({
            'title': title,
            'code': code.strip()
        })
    
    return diagrams

def render_with_kroki(mermaid_code):
    """Render diagram using Kroki service (more reliable than mermaid.ink)"""
    try:
        url = "https://kroki.io/mermaid/png"
        headers = {'Content-Type': 'text/plain'}
        
        response = requests.post(url, data=mermaid_code.encode('utf-8'), headers=headers, timeout=30)
        
        if response.status_code == 200:
            return BytesIO(response.content)
        else:
            print(f"      Error: Status {response.status_code}")
            return None
    except Exception as e:
        print(f"      Error: {e}")
        return None

def render_with_quickchart(mermaid_code):
    """Render using QuickChart API (another alternative)"""
    try:
        import urllib.parse
        encoded = urllib.parse.quote(mermaid_code)
        url = f"https://quickchart.io/chart?c={encoded}&cht=mermaid&w=1200&h=800"
        
        response = requests.get(url, timeout=30)
        
        if response.status_code == 200:
            return BytesIO(response.content)
        else:
            return None
    except:
        return None

def insert_diagrams_into_word(doc_path, diagrams, output_path):
    """Insert rendered diagrams into Word document"""
    print(f"\nProcessing: {doc_path}")
    
    try:
        doc = Document(doc_path)
    except Exception as e:
        print(f"  X Error opening document: {e}")
        return False
    
    # Add appendix section
    doc.add_page_break()
    doc.add_paragraph('Appendix A: System Architecture Diagrams', style='Heading 1')
    doc.add_paragraph('The following diagrams illustrate the system architecture, workflows, and technical design.')
    doc.add_paragraph()
    
    success_count = 0
    for i, diagram in enumerate(diagrams, 1):
        print(f"  [{i}/{len(diagrams)}] {diagram['title']}...")
        
        # Add title
        doc.add_paragraph(diagram['title'], style='Heading 2')
        
        # Try Kroki first
        image_data = render_with_kroki(diagram['code'])
        
        # If Kroki fails, try QuickChart
        if not image_data:
            print(f"      Trying alternative service...")
            image_data = render_with_quickchart(diagram['code'])
        
        if image_data:
            try:
                doc.add_picture(image_data, width=Inches(6.5))
                doc.add_paragraph()
                print(f"      > Inserted successfully")
                success_count += 1
            except Exception as e:
                print(f"      X Failed to insert: {e}")
                doc.add_paragraph("[Diagram rendering failed - see DIAGRAMS_MERMAID.md]")
        else:
            print(f"      X Failed to render")
            doc.add_paragraph("[Diagram could not be rendered - see DIAGRAMS_MERMAID.md for code]")
            doc.add_paragraph()
    
    # Save
    try:
        doc.save(output_path)
        print(f"  > Saved: {output_path} ({success_count}/{len(diagrams)} diagrams inserted)")
        return success_count > 0
    except Exception as e:
        print(f"  X Error saving: {e}")
        return False

def main():
    print("=" * 70)
    print("DIAGRAM EXPORT FIX - Using Kroki API")
    print("=" * 70)
    print("\nUsing Kroki.io - a more reliable diagram rendering service")
    print("=" * 70)
    
    # Extract diagrams
    print("\n[STEP 1] Extracting diagrams...")
    try:
        diagrams = extract_mermaid_diagrams('DIAGRAMS_MERMAID.md')
        print(f"   > Found {len(diagrams)} diagrams")
    except Exception as e:
        print(f"   X Error: {e}")
        return
    
    # Process documents
    print("\n[STEP 2] Processing Word documents...")
    
    documents = [
        ('Hospital_Management_System_PMP_with_Budget.docx', 
         'Hospital_Management_System_PMP_with_Budget_FINAL_v2.docx',
         'Arabic Version'),
        ('Hospital_Management_System_PMP_English.docx',
         'Hospital_Management_System_PMP_English_FINAL_v2.docx',
         'English Version')
    ]
    
    processed = 0
    for input_doc, output_doc, description in documents:
        if os.path.exists(input_doc):
            print(f"\n{description}:")
            if insert_diagrams_into_word(input_doc, diagrams, output_doc):
                processed += 1
        else:
            print(f"\n{description}: X Not found ({input_doc})")
    
    # Summary
    print("\n" + "=" * 70)
    if processed > 0:
        print("SUCCESS!")
        print("=" * 70)
        print(f"\n{processed} document(s) created with diagrams:")
        for _, output_doc, desc in documents:
            if os.path.exists(output_doc):
                print(f"  - {output_doc}")
        print("\nNew documents created with '_v2' suffix")
    else:
        print("FAILED - No diagrams rendered")
        print("=" * 70)
    
    print()

if __name__ == '__main__':
    try:
        import requests
        from docx import Document
        main()
    except ImportError as e:
        print(f"Missing package: {e}")
        print("Install with: pip install python-docx requests")

