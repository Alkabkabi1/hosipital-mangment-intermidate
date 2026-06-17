"""
Automated Mermaid Diagram Exporter and Word Document Inserter
This script will:
1. Extract all Mermaid diagrams from DIAGRAMS_MERMAID.md
2. Render them as PNG images using Mermaid CLI
3. Insert them into the Word documents

Requirements:
    npm install -g @mermaid-js/mermaid-cli
    pip install python-docx
"""

import subprocess
import os
import re
from docx import Document
from docx.shared import Inches

def extract_mermaid_diagrams(md_file):
    """Extract all mermaid diagrams from markdown file"""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Find all diagram sections
    pattern = r'## (Diagram \d+: [^\n]+)\n\n```mermaid\n(.*?)```'
    matches = re.findall(pattern, content, re.DOTALL)
    
    diagrams = []
    for title, code in matches:
        diagrams.append({
            'title': title,
            'code': code.strip()
        })
    
    return diagrams

def save_mermaid_files(diagrams, output_dir='temp_diagrams'):
    """Save each diagram to a separate .mmd file"""
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    files = []
    for i, diagram in enumerate(diagrams, 1):
        filename = f'{output_dir}/diagram_{i}.mmd'
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(diagram['code'])
        files.append({
            'mmd_file': filename,
            'png_file': f'{output_dir}/diagram_{i}.png',
            'title': diagram['title']
        })
    
    return files

def render_diagrams_with_mmdc(diagram_files):
    """Render diagrams using mermaid CLI (mmdc)"""
    print("Rendering diagrams with Mermaid CLI...")
    
    for diagram in diagram_files:
        cmd = [
            'mmdc',
            '-i', diagram['mmd_file'],
            '-o', diagram['png_file'],
            '-b', 'white',
            '-w', '1200',
            '-H', '800'
        ]
        
        try:
            subprocess.run(cmd, check=True, capture_output=True)
            print(f"✓ Rendered: {diagram['title']}")
        except subprocess.CalledProcessError as e:
            print(f"✗ Failed to render: {diagram['title']}")
            print(f"  Error: {e.stderr.decode()}")
        except FileNotFoundError:
            print("\n❌ ERROR: Mermaid CLI (mmdc) not found!")
            print("Please install it with: npm install -g @mermaid-js/mermaid-cli")
            return False
    
    return True

def insert_diagrams_into_word(doc_path, diagram_files, output_path):
    """Insert rendered diagrams into Word document appendix"""
    try:
        doc = Document(doc_path)
    except Exception as e:
        print(f"Error opening document: {e}")
        return False
    
    # Find the appendix section
    appendix_found = False
    insert_index = None
    
    for i, paragraph in enumerate(doc.paragraphs):
        # Look for "Appendix A" or similar
        if 'Appendix' in paragraph.text and ('System' in paragraph.text or 'Diagram' in paragraph.text):
            appendix_found = True
            insert_index = i + 1
            break
    
    if not appendix_found:
        # If no appendix found, add at the end
        doc.add_page_break()
        doc.add_paragraph('Appendix A: System Diagrams', style='Heading 1')
        insert_index = len(doc.paragraphs)
    
    # Insert diagrams
    for diagram in diagram_files:
        if os.path.exists(diagram['png_file']):
            # Add diagram title
            doc.add_paragraph(diagram['title'], style='Heading 2')
            
            # Add the image
            try:
                doc.add_picture(diagram['png_file'], width=Inches(6.5))
                doc.add_paragraph()  # Add spacing
                print(f"✓ Inserted: {diagram['title']}")
            except Exception as e:
                print(f"✗ Failed to insert: {diagram['title']} - {e}")
        else:
            print(f"✗ Image not found: {diagram['png_file']}")
    
    # Save the document
    doc.save(output_path)
    print(f"\n✓ Document saved: {output_path}")
    return True

def main():
    print("=" * 70)
    print("Automated Mermaid Diagram to Word Document Inserter")
    print("=" * 70)
    print()
    
    # Step 1: Extract diagrams
    print("Step 1: Extracting diagrams from DIAGRAMS_MERMAID.md...")
    diagrams = extract_mermaid_diagrams('DIAGRAMS_MERMAID.md')
    print(f"✓ Found {len(diagrams)} diagrams")
    print()
    
    # Step 2: Save to individual files
    print("Step 2: Creating temporary diagram files...")
    diagram_files = save_mermaid_files(diagrams)
    print(f"✓ Created {len(diagram_files)} .mmd files")
    print()
    
    # Step 3: Render to PNG
    print("Step 3: Rendering diagrams to PNG...")
    if not render_diagrams_with_mmdc(diagram_files):
        print("\n⚠ Skipping rendering. You'll need to render manually.")
        print("\nAlternative: Use https://mermaid.live to export each diagram manually")
        return
    print()
    
    # Step 4: Insert into Word documents
    print("Step 4: Inserting diagrams into Word documents...")
    
    # Arabic version
    if os.path.exists('Hospital_Management_System_PMP_with_Budget.docx'):
        print("\nProcessing Arabic version...")
        insert_diagrams_into_word(
            'Hospital_Management_System_PMP_with_Budget.docx',
            diagram_files,
            'Hospital_Management_System_PMP_with_Budget_FINAL.docx'
        )
    
    # English version
    if os.path.exists('Hospital_Management_System_PMP_English.docx'):
        print("\nProcessing English version...")
        insert_diagrams_into_word(
            'Hospital_Management_System_PMP_English.docx',
            diagram_files,
            'Hospital_Management_System_PMP_English_FINAL.docx'
        )
    
    print()
    print("=" * 70)
    print("✓ COMPLETE!")
    print("=" * 70)
    print("\nNew documents created:")
    print("  • Hospital_Management_System_PMP_with_Budget_FINAL.docx")
    print("  • Hospital_Management_System_PMP_English_FINAL.docx")
    print("\nYou can delete the 'temp_diagrams' folder after reviewing the documents.")

if __name__ == '__main__':
    main()

